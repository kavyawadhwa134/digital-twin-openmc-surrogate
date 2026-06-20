from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kavyawadhwa/Documents/Digital Twin")
OUT_DIR = ROOT / "outputs"
FIG_DIR = ROOT / "figures"
MODEL_DIR = ROOT / "models"
DATA_DIR = ROOT / "data" / "processed"

DOCX_OUT = OUT_DIR / "ICTP_Digital_Twin_Final_Results.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 40)
MUTED = RGBColor(90, 98, 108)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"
WHITE = "FFFFFF"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text())


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="") as fp:
        return list(csv.DictReader(fp))


def mean_column(rows: list[dict[str, str]], column: str) -> float:
    values = [float(row[column]) for row in rows if row.get(column) not in (None, "")]
    return sum(values) / len(values)


def fmt_float(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{100.0 * value:.{digits}f}%"


def fmt_pct_raw(value: float, digits: int = 1) -> str:
    return f"{value:.{digits}f}%"


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths_dxa: list[int], indent_dxa: int = 120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = col_widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, col_widths_dxa: list[int], header_fill=LIGHT_BLUE, body_fill=WHITE):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table, 9360, col_widths_dxa)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            set_cell_shading(cell, header_fill if r_idx == 0 else body_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.0 if r_idx else 9.2,
                        bold=True if r_idx == 0 else None,
                        color=INK,
                    )
                if c_idx > 0 and r_idx > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_doc_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, margin, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def set_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    left = hp.add_run("ICTP Digital Twin Result Brief")
    set_run_font(left, size=9, color=MUTED, bold=True)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    run = fp.add_run("Generated from local OpenMC/ML result artifacts | June 15, 2026")
    set_run_font(run, size=8.5, color=MUTED)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("RESULT BRIEF")
    set_run_font(run, size=11, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Physics-Informed Digital Twin and OpenMC Surrogate Results")
    set_run_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "OpenMC pin-cell response surrogate and multi-horizon reactor-state anomaly forecaster"
    )
    set_run_font(run, size=12.5, color=MUTED)

    rows = [
        ("Project", "Machine-learned surrogate modelling for digital twin applications"),
        ("Reactor scope", "Generation IV-relevant pin-cell / material-kernel proof of concept"),
        ("Data sources", "OpenMC high-stat and broad pin-cell sweeps, physics-informed state trajectories"),
        ("Prepared for", "ICTP poster discussion"),
        ("Generated", date(2026, 6, 15).strftime("%B %d, %Y")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for r, (label, value) in enumerate(rows):
        table.cell(r, 0).text = label
        table.cell(r, 1).text = value
    style_table(table, [2100, 7260], header_fill=LIGHT_GRAY, body_fill=WHITE)

    add_note_box(
        doc,
        "Headline interpretation",
        "Within the sampled pin-cell and simulated state-trajectory domain, the workflow "
        "demonstrates fast surrogate response prediction and early anomaly detection. "
        "It should be presented as physics-informed simulated validation, not full-core or real-plant validation.",
    )


def add_note_box(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell, color="D0D7E2", size="6")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, size=10.5, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.2, color=INK)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color=MUTED)


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.2):
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"Missing figure: {path}")
        set_run_font(r, size=10, bold=True, color=RGBColor(155, 28, 28))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def page_break(doc: Document):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    pincell = load_json(MODEL_DIR / "pincell_lhs120_highstat_engineered_surrogate_metrics.json")
    pincell_broad = load_json(MODEL_DIR / "pincell_lhs500_engineered_surrogate_metrics.json")
    pincell_rows = load_csv(DATA_DIR / "pincell_lhs120_highstat_openmc.csv")
    pincell_broad_rows = load_csv(DATA_DIR / "pincell_lhs500_openmc.csv")
    h10 = load_json(MODEL_DIR / "state_forecaster_h10_metrics.json")
    h10_val = load_csv(DATA_DIR / "state_forecaster_h10_validation_summary.csv")
    robustness = load_json(MODEL_DIR / "digital_twin_robustness_metrics.json")
    horizon_rows = robustness["horizons"]
    robustness_rows = robustness["robustness"]
    detection_rows = robustness["anomaly_detection"]
    classification_rows = robustness["fault_classification"]

    def horizon_row(seconds: int) -> dict:
        return next(row for row in horizon_rows if int(row["horizon_seconds"]) == seconds)

    def robustness_row(seconds: int, condition: str) -> dict:
        return next(
            row
            for row in robustness_rows
            if int(row["horizon_seconds"]) == seconds and row["condition"] == condition
        )

    def detection_row(seconds: int, split: str) -> dict:
        return next(
            row
            for row in detection_rows
            if int(row["horizon_seconds"]) == seconds and row["split"] == split
        )

    def classification_row(seconds: int) -> dict:
        return next(row for row in classification_rows if int(row["horizon_seconds"]) == seconds)

    doc = Document()
    set_doc_styles(doc)
    set_header_footer(doc)
    add_title_page(doc)

    doc.add_heading("1. Executive Result Summary", level=1)
    summary_rows = [
        [
            "Pin-cell response surrogate",
            "OpenMC-labeled reactor response prediction",
            f"keff MAE {pincell['summary']['keff_mae_pcm']:.0f} pcm; "
            f"{pincell['summary']['keff_percent_within_1000_pcm']:.0f}% within 1000 pcm; "
            f"mean target relative MAE {100.0 * pincell['summary']['mean_relative_mae_across_targets']:.2f}%; "
            f"about {pincell['speedup_vs_openmc_pincell']:,.0f}x apparent response-evaluation speedup",
        ],
        [
            "10 s digital-twin forecaster",
            "Predict normal reactor state and detect deviations",
            f"+10 s endpoint fuel T RMSE {h10['endpoint_feature_metrics']['fuel_temperature_K']['rmse']:.2f} K; "
            f"keff RMSE {h10['endpoint_feature_metrics']['keff']['rmse'] * 1.0e5:.1f} pcm; "
            f"known-fault F1 {detection_row(10, 'known_fault_families')['f1']:.3f}; "
            f"1x sensor-noise alarm rate {100.0 * robustness_row(10, 'gaussian_sensor_noise_1x')['normal_alarm_rate']:.2f}%",
        ],
        [
            "OpenMC/ML cross-section visuals",
            "Geometry and response-surface evidence",
            "OpenMC pin-cell material map plus ML keff/flux response surfaces; microscopic ENDF XS surrogate removed from headline claims",
        ],
    ]
    add_table_from_rows(
        doc,
        ["Component", "Role", "Main result"],
        summary_rows,
        [2300, 2450, 4610],
    )
    add_note_box(
        doc,
        "Use this wording on the poster",
        "The pin-cell surrogate is a fast response surrogate inside the sampled OpenMC domain. "
        "The digital twin forecasts normal reactor evolution and uses forecast residuals plus physics residuals for anomaly detection. "
        "The microscopic cross-section surrogate should be treated as future XSBench/GPU work, not as a current speed result.",
    )

    doc.add_heading("2. Model 1: OpenMC Pin-Cell Response Surrogate", level=1)
    doc.add_paragraph(
        "This model maps operating and geometry parameters to OpenMC-computed response quantities: "
        "keff, fuel/moderator fluxes, fission rate, capture rates, and a power-density proxy. "
        "The revised headline uses the 120-case high-stat Latin-hypercube dataset, not the larger 500-case low-stat sweep, "
        "because keff label noise is the limiting error source."
    )
    pc_summary = pincell["summary"]
    pc_report = pincell["validation_report"]
    add_table_from_rows(
        doc,
        ["Metric", "Random interpolation", "Hot-fuel holdout", "Wide-pitch holdout"],
        [
            [
                "keff MAE",
                f"{pc_report['random_interpolation']['keff']['mae_pcm']:.0f} pcm",
                f"{pc_report['held_out_hot_fuel_regime']['keff']['mae_pcm']:.0f} pcm",
                f"{pc_report['held_out_wide_pitch_geometry']['keff']['mae_pcm']:.0f} pcm",
            ],
            [
                "keff RMSE",
                f"{pc_report['random_interpolation']['keff']['rmse_pcm']:.0f} pcm",
                f"{pc_report['held_out_hot_fuel_regime']['keff']['rmse_pcm']:.0f} pcm",
                f"{pc_report['held_out_wide_pitch_geometry']['keff']['rmse_pcm']:.0f} pcm",
            ],
            [
                "Within 500 pcm",
                fmt_pct_raw(pc_report["random_interpolation"]["keff"]["percent_within_500_pcm"]),
                fmt_pct_raw(pc_report["held_out_hot_fuel_regime"]["keff"]["percent_within_500_pcm"]),
                fmt_pct_raw(pc_report["held_out_wide_pitch_geometry"]["keff"]["percent_within_500_pcm"]),
            ],
            [
                "Within 1000 pcm",
                fmt_pct_raw(pc_report["random_interpolation"]["keff"]["percent_within_1000_pcm"]),
                fmt_pct_raw(pc_report["held_out_hot_fuel_regime"]["keff"]["percent_within_1000_pcm"]),
                fmt_pct_raw(pc_report["held_out_wide_pitch_geometry"]["keff"]["percent_within_1000_pcm"]),
            ],
            [
                "Mean relative MAE across targets",
                fmt_pct(pc_summary["mean_relative_mae_across_targets"]),
                "Reported per target in metrics JSON",
                "Reported per target in metrics JSON",
            ],
            [
                "Inference speed",
                f"{pincell['ml_predict_seconds_per_case']:.2e} s/case",
                f"OpenMC mean {pincell['openmc_mean_elapsed_seconds_per_case']:.2f} s/case",
                f"Apparent speedup {pincell['speedup_vs_openmc_pincell']:,.0f}x",
            ],
        ],
        [2450, 2300, 2300, 2310],
    )
    add_figure(
        doc,
        FIG_DIR / "pincell_lhs120_highstat_engineered_surrogate_vs_openmc.png",
        "Figure 1. keff prediction versus OpenMC for the high-stat pin-cell surrogate.",
        5.0,
    )
    broad_report = pincell_broad["validation_report"]
    add_table_from_rows(
        doc,
        ["Dataset", "OpenMC settings", "Mean keff std", "Random keff MAE", "Within 1000 pcm", "Conclusion"],
        [
            [
                "120 high-stat LHS",
                f"{int(float(pincell_rows[0]['batches']))} batches, {int(float(pincell_rows[0]['particles'])):,} particles",
                f"{mean_column(pincell_rows, 'keff_std') * 1.0e5:.0f} pcm",
                f"{pc_report['random_interpolation']['keff']['mae_pcm']:.0f} pcm",
                fmt_pct_raw(pc_report["random_interpolation"]["keff"]["percent_within_1000_pcm"]),
                "Best headline model",
            ],
            [
                "500 broad LHS",
                f"{int(float(pincell_broad_rows[0]['batches']))} batches, {int(float(pincell_broad_rows[0]['particles'])):,} particles",
                f"{mean_column(pincell_broad_rows, 'keff_std') * 1.0e5:.0f} pcm",
                f"{broad_report['random_interpolation']['keff']['mae_pcm']:.0f} pcm",
                fmt_pct_raw(broad_report["random_interpolation"]["keff"]["percent_within_1000_pcm"]),
                "Useful coverage, noisier labels",
            ],
        ],
        [1500, 1700, 1350, 1500, 1450, 1860],
    )
    add_note_box(
        doc,
        "Why this is the stronger pin-cell result",
        "The project now treats OpenMC label quality as part of the surrogate design. "
        "The broader 500-case sweep is useful for coverage, but its lower particle count increases label uncertainty and worsens keff error. "
        "For the poster, the defensible claim is a high-stat, uncertainty-aware response surrogate within the sampled pin-cell domain.",
    )
    add_figure(
        doc,
        FIG_DIR / "pincell_surrogate_highstat_vs_largedata.png",
        "Figure 2. High-stat OpenMC labels produce a stronger keff surrogate than a larger but noisier low-stat sweep.",
        6.2,
    )
    add_figure(
        doc,
        FIG_DIR / "pincell_lhs120_highstat_engineered_multioutput_surrogate_vs_openmc.png",
        "Figure 3. Multi-output OpenMC response surrogate comparison across flux, rates, and power proxy.",
        6.3,
    )

    page_break(doc)
    doc.add_heading("3. Model 2: 10-Second Digital-Twin State Forecaster", level=1)
    doc.add_paragraph(
        "This model uses the previous 20 seconds of state history to forecast the true +10 second endpoint state. "
        "The selected forecaster is an ensemble of regularized, balanced, and wide MLP predictors, chosen by a validation score weighted toward endpoint accuracy. "
        "The anomaly detector combines forecast residuals, signed residual-growth features, and physics-consistency residuals. "
        "Detector calibration now includes noisy, delayed, and dropout-corrupted normal windows so routine sensor imperfections are less likely to become false alarms."
    )
    endpoint = h10["endpoint_feature_metrics"]
    add_table_from_rows(
        doc,
        ["Predicted quantity", "+10 s endpoint RMSE", "+10 s endpoint MAE", "Interpretation"],
        [
            ["Fuel temperature", f"{endpoint['fuel_temperature_K']['rmse']:.2f} K", f"{endpoint['fuel_temperature_K']['mae']:.2f} K", "Temperature trajectory remains tight in normal simulated dynamics"],
            ["Coolant outlet temperature", f"{endpoint['coolant_temperature_K']['rmse']:.3f} K", f"{endpoint['coolant_temperature_K']['mae']:.3f} K", "Small thermal-state error"],
            ["keff", f"{endpoint['keff']['rmse'] * 1e5:.1f} pcm", f"{endpoint['keff']['mae'] * 1e5:.1f} pcm", "Low reactivity-state forecast error"],
            ["Reactor power", f"{endpoint['power_norm']['rmse']:.4f}", f"{endpoint['power_norm']['mae']:.4f}", "About 0.49% normalized RMSE"],
            ["Fast flux", f"{endpoint['flux_fast_norm']['rmse']:.4f}", f"{endpoint['flux_fast_norm']['mae']:.4f}", "About 0.34% normalized RMSE"],
            ["Fission rate", f"{endpoint['fission_rate_norm']['rmse']:.4f}", f"{endpoint['fission_rate_norm']['mae']:.4f}", "About 0.51% normalized RMSE"],
            ["Capture rate", f"{endpoint['capture_rate_norm']['rmse']:.4f}", f"{endpoint['capture_rate_norm']['mae']:.4f}", "About 0.28% normalized RMSE"],
        ],
        [2200, 1900, 1900, 3360],
    )
    doc.add_heading("Anomaly validation", level=2)
    split_names = {
        "mixed_heldout": "Mixed held-out faults",
        "unseen_family": "Unseen fault families",
        "weak_transient_stress": "Weak-transient stress",
    }
    add_table_from_rows(
        doc,
        ["Validation split", "Precision", "Recall", "F1", "False-positive rate", "Event alarm delay"],
        [
            [
                split_names.get(row["split"], row["split"]),
                fmt_float(float(row["precision"])),
                fmt_float(float(row["recall"])),
                fmt_float(float(row["f1"])),
                fmt_pct(float(row["false_positive_rate"]), 3),
                f"median {float(row['median_alarm_delay_seconds']):.0f} s; p90 {float(row['p90_alarm_delay_seconds']):.1f} s",
            ]
            for row in h10_val
        ],
        [2250, 1250, 1250, 1250, 1650, 1710],
    )
    add_note_box(
        doc,
        "Interpretation of anomaly behavior",
        "The forecaster predicts normal reactor evolution. After a fault begins, disagreement between predicted normal evolution and observed state is expected; "
        "that residual is the anomaly signal. The reported alarm occurs at 282 s for an injected anomaly beginning at 280 s.",
    )
    add_figure(
        doc,
        FIG_DIR / "state_forecaster_h10_anomaly_detection.png",
        "Figure 4. True +10 s endpoint forecasts for flux, power, coolant temperature, fuel temperature, and reactivity, with anomaly alarm probability.",
        6.3,
    )
    add_figure(
        doc,
        FIG_DIR / "state_forecaster_h10_validation_summary.png",
        "Figure 5. Window-level anomaly detection scores and event-level alarm timing for held-out validation splits.",
        5.9,
    )

    page_break(doc)
    doc.add_heading("4. OpenMC Geometry and ML Response Cross-Section Images", level=1)
    doc.add_paragraph(
        "The cross-section figures below are geometry and response-surface cross sections of the simulated pin-cell problem. "
        "They should not be confused with a validated microscopic ENDF cross-section surrogate."
    )
    add_note_box(
        doc,
        "Removed weak XS speed claim",
        "The separate ENDF microscopic cross-section surrogate was removed from the headline results because the current CPU sklearn implementation is slower than OpenMC's vectorized lookup. "
        "A future claim should be rebuilt around XSBench-style kernels, batched neural inference, and Apple Metal/GPU or MPI benchmarking.",
    )
    add_figure(
        doc,
        FIG_DIR / "openmc_pincell_geometry_cross_section.png",
        "Figure 6. OpenMC pin-cell material geometry cross-section used for the response-surrogate demonstration.",
        4.7,
    )
    add_figure(
        doc,
        FIG_DIR / "ml_surrogate_response_cross_section.png",
        "Figure 7. ML surrogate response surface over the pin-cell parameter cross-section.",
        5.6,
    )
    add_figure(
        doc,
        FIG_DIR / "openmc_vs_ml_cross_section_summary.png",
        "Figure 8. Combined OpenMC geometry and ML response cross-section summary.",
        5.8,
    )

    page_break(doc)
    doc.add_heading("5. Robustness and Presentation Boundaries", level=1)
    h10_robust = horizon_row(10)
    h20_robust = horizon_row(20)
    h30_robust = horizon_row(30)
    h60_robust = horizon_row(60)
    add_table_from_rows(
        doc,
        ["Forecast horizon", "Fuel T RMSE", "keff RMSE", "Power RMSE", "90% interval coverage", "Clean false alarms"],
        [
            [
                "+10 s",
                f"{h10_robust['fuel_temperature_endpoint_rmse_K']:.2f} K",
                f"{h10_robust['keff_endpoint_rmse_pcm']:.1f} pcm",
                f"{h10_robust['power_endpoint_rmse_pct']:.2f}%",
                fmt_pct(h10_robust["mean_90_coverage"]),
                fmt_pct(h10_robust["long_normal_false_alarm_rate"], 3),
            ],
            [
                "+20 s",
                f"{h20_robust['fuel_temperature_endpoint_rmse_K']:.2f} K",
                f"{h20_robust['keff_endpoint_rmse_pcm']:.1f} pcm",
                f"{h20_robust['power_endpoint_rmse_pct']:.2f}%",
                fmt_pct(h20_robust["mean_90_coverage"]),
                fmt_pct(h20_robust["long_normal_false_alarm_rate"], 3),
            ],
            [
                "+30 s",
                f"{h30_robust['fuel_temperature_endpoint_rmse_K']:.2f} K",
                f"{h30_robust['keff_endpoint_rmse_pcm']:.1f} pcm",
                f"{h30_robust['power_endpoint_rmse_pct']:.2f}%",
                fmt_pct(h30_robust["mean_90_coverage"]),
                fmt_pct(h30_robust["long_normal_false_alarm_rate"], 3),
            ],
            [
                "+60 s",
                f"{h60_robust['fuel_temperature_endpoint_rmse_K']:.2f} K",
                f"{h60_robust['keff_endpoint_rmse_pcm']:.1f} pcm",
                f"{h60_robust['power_endpoint_rmse_pct']:.2f}%",
                fmt_pct(h60_robust["mean_90_coverage"]),
                fmt_pct(h60_robust["long_normal_false_alarm_rate"], 3),
            ],
        ],
        [1500, 1500, 1450, 1450, 1900, 1560],
    )
    add_figure(
        doc,
        FIG_DIR / "digital_twin_horizon_degradation.png",
        "Figure 9. Forecast endpoint accuracy across +10, +20, +30, and +60 second horizons.",
        5.9,
    )
    add_table_from_rows(
        doc,
        ["Robustness test", "+10 s result", "Interpretation"],
        [
            [
                "1x sensor noise",
                f"Endpoint error x{robustness_row(10, 'gaussian_sensor_noise_1x')['rmse_multiplier_vs_clean']:.2f}; "
                f"alarm rate {100.0 * robustness_row(10, 'gaussian_sensor_noise_1x')['normal_alarm_rate']:.2f}%",
                "Routine sensor noise mostly affects uncertainty/error, not alarm state.",
            ],
            [
                "3 s signal delay",
                f"Endpoint error x{robustness_row(10, 'three_second_signal_delay')['rmse_multiplier_vs_clean']:.2f}; "
                f"alarm rate {100.0 * robustness_row(10, 'three_second_signal_delay')['normal_alarm_rate']:.2f}%",
                "Short delays are handled well after robust calibration.",
            ],
            [
                "8% sensor dropout",
                f"Endpoint error x{robustness_row(10, 'eight_percent_dropout_hold_last')['rmse_multiplier_vs_clean']:.2f}; "
                f"alarm rate {100.0 * robustness_row(10, 'eight_percent_dropout_hold_last')['normal_alarm_rate']:.2f}%",
                "Hold-last-value dropout imputation has negligible effect in this simulator.",
            ],
            [
                "2x noise / biased sensors",
                f"2x noise alarm rate {100.0 * robustness_row(10, 'gaussian_sensor_noise_2x')['normal_alarm_rate']:.1f}%; "
                f"bias alarm rate {100.0 * robustness_row(10, 'small_sensor_bias')['normal_alarm_rate']:.1f}%",
                "Severe noise and persistent sensor bias are treated as operational anomalies.",
            ],
            [
                "OOD operating regimes",
                f"High-power error x{robustness_row(10, 'ood_high_power_maneuver')['rmse_multiplier_vs_clean']:.1f}; "
                f"low-density error x{robustness_row(10, 'ood_low_coolant_density')['rmse_multiplier_vs_clean']:.1f}",
                "The model detects a domain boundary; this should be presented as stress testing, not validated plant coverage.",
            ],
        ],
        [2250, 3150, 3960],
    )
    add_figure(
        doc,
        FIG_DIR / "digital_twin_robustness_summary.png",
        "Figure 10. Sensor-stress and synthetic OOD robustness summary for the +10 s digital-twin model.",
        6.2,
    )
    add_table_from_rows(
        doc,
        ["Capability", "Result", "Boundary"],
        [
            [
                "Known fault classification",
                f"+10 s accuracy {classification_row(10)['known_fault_classifier_accuracy']:.3f}; "
                f"macro-F1 {classification_row(10)['known_fault_classifier_macro_f1']:.3f}",
                "Classes are coolant loss, control-rod withdrawal, and flux-detector bias.",
            ],
            [
                "Unknown family rejection",
                f"{100.0 * classification_row(10)['unknown_family_rejection_rate']:.1f}% of coolant-heating/sensor-drift windows rejected as unknown",
                "This is a simulated unknown-family test, not real event diagnosis.",
            ],
            [
                "Known-family anomaly detection",
                f"F1 {detection_row(10, 'known_fault_families')['f1']:.3f}; "
                f"detection rate {100.0 * detection_row(10, 'known_fault_families')['detection_rate']:.1f}%",
                "Independent seeds/severities; same fault families as calibration.",
            ],
            [
                "Unseen-family anomaly detection",
                f"F1 {detection_row(10, 'unseen_fault_families')['f1']:.3f}; "
                f"detection rate {100.0 * detection_row(10, 'unseen_fault_families')['detection_rate']:.1f}%",
                "Harder because coolant heating and sensor drift are not detector-training families.",
            ],
            [
                "Physics projection",
                f"Mean physics residual improves from {h10_robust['prediction_physics_residual_mean_abs_z']:.3f} to "
                f"{h10_robust['projected_prediction_physics_residual_mean_abs_z']:.3f} |z|",
                "Implemented as post-prediction projection; fully constrained training remains future work.",
            ],
        ],
        [2450, 3500, 3410],
    )
    add_figure(
        doc,
        FIG_DIR / "digital_twin_uncertainty_classification.png",
        "Figure 11. Forecast uncertainty coverage, physics residual behavior, and +10 s known-fault classifier confusion matrix.",
        6.3,
    )
    add_note_box(
        doc,
        "Most defensible final claim",
        "The project demonstrates a physics-informed surrogate and digital-twin workflow for Generation IV-relevant reactor kernels. "
        "The pin-cell model provides fast response prediction within the sampled OpenMC domain; the state forecaster predicts normal +10 to +60 s reactor evolution and detects simulated anomalies early. "
        "The strongest claim is simulated, physics-informed digital-twin validation with independent seeds, sensor-stress testing, uncertainty intervals, and fault-family classification. "
        "Full-core validation, real plant data, strictly constrained neural training, and microscopic XSBench/GPU kernel replacement remain future work.",
    )

    doc.save(DOCX_OUT)
    return DOCX_OUT


if __name__ == "__main__":
    print(build_doc())
